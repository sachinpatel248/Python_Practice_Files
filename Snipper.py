from Tkinter import *
import Tkinter, tkFileDialog
import glob
import docx
import os
import pyHook
import pygame
from PIL import Image, ImageTk
from PIL import ImageGrab
import time
import threading
import ctypes
import re
from win32api import *
from win32gui import *
import win32con
import sys, os
import struct
import pyautogui



ImageEditorPath=r'C:\WINDOWS\system32\mspaint.exe'

threads = []
workingDir= os.getcwd()
Iconpath=os.getcwd()+str(r'\SnipperIcon.ico')
BgImagePath=os.getcwd()+str(r'\bg.jpg')


master = Tk()
master.wm_title("SNIPPER")
master.geometry("500x300")
width = 500 # width for the Tk root
height = 300 # height for the Tk root

# get screen width and height
widthscreen = master.winfo_screenwidth() # width of the screen
heightscreen = master.winfo_screenheight() # height of the screen

# calculate x and y coordinates for the Tk root window
x = (widthscreen/2) - (width/2)
y = (heightscreen/2) - (height/2)

# set the dimensions of the screen 
# and where it is placed
master.geometry('%dx%d+%d+%d' % (width, height, x, y))
master.resizable(width=FALSE, height=FALSE)
startStartedFlag=False
wantToOpenImageInEditor=False

def Notify(Message):
    class WindowsBalloonTip:
        def __init__(self, title, msg):
            message_map = {
                    win32con.WM_DESTROY: self.OnDestroy,
            }
            # Register the Window class.
            wc = WNDCLASS()
            hinst = wc.hInstance = GetModuleHandle(None)
            wc.lpszClassName = "PythonTaskbar"
            wc.lpfnWndProc = message_map # could also specify a wndproc.
            classAtom = RegisterClass(wc)
            # Create the Window.
            style = win32con.WS_OVERLAPPED | win32con.WS_SYSMENU
            self.hwnd = CreateWindow( classAtom, "Taskbar", style, \
                    0, 0, win32con.CW_USEDEFAULT, win32con.CW_USEDEFAULT, \
                    0, 0, hinst, None)
            UpdateWindow(self.hwnd)
            iconPathName = os.path.abspath(os.path.join( sys.path[0], Iconpath ))
            icon_flags = win32con.LR_LOADFROMFILE | win32con.LR_DEFAULTSIZE
            try:
               hicon = LoadImage(hinst, iconPathName, \
                        win32con.IMAGE_ICON, 0, 0, icon_flags)
            except:
              hicon = LoadIcon(0, win32con.IDI_APPLICATION)
            flags = NIF_ICON | NIF_MESSAGE | NIF_TIP
            nid = (self.hwnd, 0, flags, win32con.WM_USER+20, hicon, "tooltip")
            Shell_NotifyIcon(NIM_ADD, nid)
            Shell_NotifyIcon(NIM_MODIFY, \
                             (self.hwnd, 0, NIF_INFO, win32con.WM_USER+20,\
                              hicon, "Balloon  tooltip",msg,200,title))
            # self.show_balloon(title, msg)
            time.sleep(1)
            DestroyWindow(self.hwnd)
            classAtom = UnregisterClass(classAtom, hinst)
        def OnDestroy(self, hwnd, msg, wparam, lparam):
            nid = (self.hwnd, 0)
            Shell_NotifyIcon(NIM_DELETE, nid)
            PostQuitMessage(0) # Terminate the app.

    def balloon_tip(title, msg):
        w=WindowsBalloonTip(title, msg)

    if __name__ == '__main__':
        balloon_tip("Snipper Notifier", Message)


     
def CreateDocx():
    if(ValidatePath()):
        doc = docx.Document()
        num=1
        path=w.get()
        path=os.path.join(os.path.abspath(path),"*.jpeg")
        for filename in glob.glob(path):
            try:
                doc.add_heading(str(num), 4).style='List'
                doc.add_picture(filename, width=docx.shared.Inches(6.5),height=docx.shared.Inches(4))
                num=num+1
            except Exception as e:
                pass
			 
        Message ="Document with name 'AutoDocument.docx' created with "+str(num)+" images in selected folder."
        path=w.get()
        path=path+r'/AutoDocument.docx'
        doc.save(path)
        Notify(Message)


def Browse():
     path = tkFileDialog.askdirectory()
     w.delete(0, END)
     w.insert(0,path)
     
def StartOnce():

	def TakeScreenShot():
		"in screnshot"
		try:
			img = ImageGrab.grab()
			SaveDirectory=w.get()

			filename=r'\ScreenShot'+str(time.strftime("%Y-%m-%d--%H-%M-%S"))+'.jpeg'
			filepath=SaveDirectory+filename
			img.save(filepath)
			Message="Image captured with name "+str(filename[1:])
			print (Message)
			t = threading.Thread(target=Notify(Message))
			threads.append(t)
			t.start()


		except Exception as e:
			print (str(e))

	def TakeScreenShotAndOpen():
		"in editor"
		try:
			global ImageEditorPath
			img = ImageGrab.grab()

			SaveDirectory=w.get()



			#print savednumber
			filename=r'\ScreenShot'+str(time.strftime("%Y-%m-%d--%H-%M-%S"))+'.jpeg'
			filepath=SaveDirectory+filename
			img.save(filepath)
			#print filepath
			
			Message="Image captured with name "+str(filename[1:])
			t = threading.Thread(target=Notify(Message))
			threads.append(t)
			t.start()

			editorstring='""%s" "%s"'% (ImageEditorPath,filepath)
			os.system(editorstring)
			


		except Exception as e:
			print (str(e))

	def OnKeyboardEvent(event):

		if event.KeyID ==45:
			t = threading.Thread(target=TakeScreenShot)
			threads.append(t)
			t.start()
			
		if event.KeyID ==44:
			t = threading.Thread(target=TakeScreenShotAndOpen)
			threads.append(t)
			t.start()
		return True


	hm = pyHook.HookManager()
	hm.KeyDown = OnKeyboardEvent
	hm.HookKeyboard()
	pygame.init()

	while True:
		pygame.event.pump()

def ValidatePath():
	path=str(w.get()).strip()
	if(path==""):
		Message="Please select folder"
		ShowMessage(Message)
		return FALSE

	if(not os.path.isdir(path)):
		Message="Please enter valid directory"
		ShowMessage(Message)
		return False

	return True
     
def ShowMessage(Message):
	global Iconpath
	root = Tk()
	root.wm_title("Snipper")
	root.geometry("350x150")
	root.resizable(width=FALSE, height=FALSE)
	root.focus_set()
	width = 300 # width for the Tk root
	height = 100 # height for the Tk root

	# get screen width and height
	widthscreen = root.winfo_screenwidth() # width of the screen
	heightscreen = root.winfo_screenheight() # height of the screen

	# calculate x and y coordinates for the Tk root window
	x = (widthscreen/2) - (width/2)
	y = (heightscreen/2) - (height/2)
	root.geometry('%dx%d+%d+%d' % (width, height, x, y))
	root.after(1, lambda: root.focus_force())
	labelInfo=Label(root, text=Message,font=("Arial",12))
	labelInfo.pack()
	labelInfo.place(x=5,y=5)
	btnOk=Button(root, text="     OK     " , command= lambda:destroyMessage())
	btnOk.place(x=200,y=60)
	btnOk.focus_force()
	global Iconpath
	root.iconbitmap(Iconpath)

	def destroyMessage():
		root.destroy()


	mainloop()

def ValidateFolderName(FolderName):
	if(FolderName==""):
		Message="Folder name should not be blank"
		ShowMessage(Message)
		return False

	if (not re.match(r'^\w+$', FolderName)):
		Message="Folder name should not contain special character"
		ShowMessage(Message)
		return False

	return True
          

               
     
          
IsFolderWindowOpened=False

def CreateNewInstance():
	global IsFolderWindowOpened
	global workingDir

	if (IsFolderWindowOpened==False):
		if(ValidatePath()):
			def CreateNewFolder(FolderName):
				if(ValidateFolderName(FolderName)):
					runDirectory=workingDir
					TochangeDir=w.get()
					os.chdir(TochangeDir)
					if (not os.path.exists(FolderName)):
						os.mkdir(FolderName)
						path=w.get()
						path=path+str('/')+str(FolderName)
						w.delete(0, END)
						w.insert(0,path)
						root.after(0, lambda: root.destroy())
						
						Message="Folder Created  "+str(FolderName)
						ShowMessage(Message)
						IsFolderWindowOpened=False
						pyautogui.hotkey('win','down')
						



					else:
						Message="Folder with sames name already exist.Try giving different folder name."
						ShowMessage(Message)
					os.chdir(runDirectory)


			IsFolderWindowOpened=True
			root = Tk()
			root.wm_title("Enter Folder Name")
			root.geometry("300x100")
			global Iconpath
			root.iconbitmap(Iconpath)
			root.resizable(width=FALSE, height=FALSE)
			root.focus_set()
			root.after(1, lambda: root.focus_force())

			def on_closing():
				 global IsFolderWindowOpened
				 IsFolderWindowOpened=False
				 root.destroy()
				 print (IsFolderWindowOpened)

			root.protocol("WM_DELETE_WINDOW", on_closing)
			width = 300 # width for the Tk root
			height = 100 # height for the Tk root

			# get screen width and height
			widthscreen = master.winfo_screenwidth() # width of the screen
			heightscreen = master.winfo_screenheight() # height of the screen

			# calculate x and y coordinates for the Tk root window
			x = (widthscreen/2) - (width/2)
			y = (heightscreen/2) - (height/2)
			root.geometry('%dx%d+%d+%d' % (width, height, x, y))
			root.geometry('%dx%d+%d+%d' % (width, height, x, y))
			labelInfo=Label(root, text="This will create Folder in working current directory")
			labelInfo.pack()
			labelInfo.place(x=5,y=5)
			LabelFolderName=Label(root, text="Folder Name : ")
			LabelFolderName.pack()
			LabelFolderName.place(x=5,y=32)
			txtFolderName = Entry(root,bd=1,width=28)
			txtFolderName.pack()
			txtFolderName.place(x=100,y=30)
			txtFolderName.focus_set()
			btnOk=Button(root, text="     OK     " , command= lambda: CreateNewFolder(txtFolderName.get()))
			btnOk.place(x=200,y=60)
			mainloop()


def NotifyToUser(Message):
	root = Tk()
	w = 350 # width for the Tk root
	h = 150 # height for the Tk root

	# get screen width and height
	ws = root.winfo_screenwidth() # width of the screen
	hs = root.winfo_screenheight() # height of the screen

	# calculate x and y coordinates for the Tk root window
	x = (ws/2) - (w/2)
	y = (hs/2) - (h/2)

	# set the dimensions of the screen 
	# and where it is placed
	root.geometry('%dx%d+%d+%d' % (w, h, x, y))
	root.wm_title("Snipper - Self destroy message")
	root.geometry("350x150")
	root.resizable(width=FALSE, height=FALSE)
	root.focus_set()
	root.after(1, lambda: root.focus_force())
	root.after(1000, lambda: root.destroy())

	labelInfo=Label(root, text=Message,font=("Arial",12))
	labelInfo.pack()
	labelInfo.place(x=5,y=5)

	btnOk=Button(root, text="     OK     " , command= lambda:destroyMessage())
	btnOk.place(x=175,y=100)
	btnOk.focus_force()
	global Iconpath
	root.iconbitmap(Iconpath)


	def destroyMessage():
		root.destroy()


	mainloop()



def Start():
	if(ValidatePath()):
		global startStartedFlag       

		if startStartedFlag==False:
			threadone= threading.Thread(None,StartOnce,None)
			threadone.start()
			startStartedFlag=True
			
			pyautogui.hotkey('win','down')
			Message="Process Started"
			NotifyToUser(Message)

		elif startStartedFlag==True:
                        
			pyautogui.hotkey('win','down')
			Message="Process already Running"
			NotifyToUser(Message)
			pyautogui.hotkey('win','down')


def KillAll():
    master.destroy()


im = Image.open(BgImagePath)
tkimage = ImageTk.PhotoImage(im)
Tkinter.Label(master,image = tkimage).pack()


defaultPath=""
w = Entry( master,bd=2,width=50)
w.pack(side = LEFT,padx=50)
w.place(x=50,y=75)
w.delete(0, END)
w.insert(0,defaultPath)



b = Button(master, text=" Browse ", command=Browse)
b.pack(side = LEFT,padx=0, pady=20)
b.place(x=400,y=75)

b0 = Button(master, text=" Start  ", command=Start)
b0.place(x=50,y=150)

b1 = Button(master, text=" Create New Instance  ", command=CreateNewInstance)
b1.place(x=175,y=150)

b2 = Button(master, text=" Create Docx" , command=CreateDocx)
b2.place(x=375,y=150)

b3 = Button(master, text="   Close all  " , command=KillAll)
b3.place(x=375,y=200)

master.iconbitmap(Iconpath)
master.focus_set()
mainloop()




